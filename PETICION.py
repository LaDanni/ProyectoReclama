import tkinter as tk
from tkinter import messagebox
import sqlite3
from docx import Document
import os

class CRUDCliente:
    def __init__(self, root):
        self.root = root
        self.root.title("CRUDCliente")
        self.root.geometry("400x500")
        self.root.resizable(True, True)

        # Variables
        self.nombre = tk.StringVar()
        self.cedula = tk.StringVar()

        # Título
        self.title_label = tk.Label(root, text="Reclamación Directa", font=('Arial', 14))
        self.title_label.grid(row=0, column=0, columnspan=2, pady=10)

        # Campos de texto
        self.create_input("Nombre", self.nombre, 1)
        self.create_input("Cédula", self.cedula, 7)

        # Botones
        self.create_button("Insertar", self.insertar, 8, 0)
        self.create_button("Editar", self.editar, 8, 1)
        self.create_button("Actualizar", self.actualizar, 9, 0)
        self.create_button("Borrar", self.borrar, 9, 1)
        self.create_button("Limpiar", self.limpiar, 10, 0, 2)
        self.create_button("Generar Documento", self.generar_documento, 11, 0, 2)

        # Lista de clientes
        self.clientes_listbox = tk.Listbox(self.root)
        self.clientes_listbox.grid(row=12, column=0, columnspan=2, padx=10, pady=10, sticky="we")
        self.root.grid_columnconfigure(0, weight=1)
        self.root.grid_columnconfigure(1, weight=1)
        self.clientes_listbox.bind("<Configure>", self.on_listbox_configure)  # Ajustar tamaño al cambiar

        # Cargar clientes
        self.cargar_clientes()

    def create_input(self, label_text, variable, row):
        label = tk.Label(self.root, text=label_text)
        label.grid(row=row, column=0, padx=10, pady=5, sticky="e")
        entry = tk.Entry(self.root, textvariable=variable)
        entry.grid(row=row, column=1, padx=10, pady=5, sticky="we")
        self.root.grid_columnconfigure(1, weight=1)

    def create_button(self, text, command, row, column, colspan=1):
        button = tk.Button(self.root, text=text, command=command)
        button.grid(row=row, column=column, padx=10, pady=5, columnspan=colspan, sticky="we")
        self.root.grid_columnconfigure(column, weight=1)

    def on_listbox_configure(self, event):
        self.clientes_listbox.config(width=event.width)

    def cargar_clientes(self):
        self.clientes_listbox.delete(0, tk.END)
        conn = sqlite3.connect('clientes.db')
        c = conn.cursor()
        c.execute("SELECT id, nombre FROM clientes")
        clientes = c.fetchall()
        for cliente in clientes:
            self.clientes_listbox.insert(tk.END, f"{cliente[0]}: {cliente[1]}")
        conn.close()

    def insertar(self):
        conn = sqlite3.connect('clientes.db')
        c = conn.cursor()
        c.execute("INSERT INTO clientes (nombre,cédula) VALUES (?, ?, ?, ?, ?, ?, ?)",
                  (self.nombre.get(), self.cedula.get()))
        conn.commit()
        conn.close()
        self.cargar_clientes()
        messagebox.showinfo("Insertar", "Cliente insertado")

    def editar(self):
        selected = self.clientes_listbox.curselection()
        if selected:
            cliente_id = self.clientes_listbox.get(selected[0]).split(":")[0]
            conn = sqlite3.connect('clientes.db')
            c = conn.cursor()
            c.execute("SELECT * FROM clientes WHERE id=?", (cliente_id,))
            cliente = c.fetchone()
            self.nombrecompleto.set(cliente[1])
            self.No. cédula.set(cliente[2])
            conn.close()
            self.editing_id = cliente_id
            messagebox.showinfo("Editar", "Cliente seleccionado para editar")

    def actualizar(self):
        if hasattr(self, 'editing_id'):
            conn = sqlite3.connect('clientes.db')
            c = conn.cursor()
            c.execute("""UPDATE clientes SET nombre=?, apellido1=?, apellido2=?, edad=?, correo=?, telefono=?, cedula=? WHERE id=?""",
                      (self.nombre.get(), self.cedula.get(), self.editing_id))
            conn.commit()
            conn.close()
            self.cargar_clientes()
            messagebox.showinfo("Actualizar", "Cliente actualizado")
            del self.editing_id

    def borrar(self):
        selected = self.clientes_listbox.curselection()
        if selected:
            cliente_id = self.clientes_listbox.get(selected[0]).split(":")[0]
            conn = sqlite3.connect('clientes.db')
            c = conn.cursor()
            c.execute("DELETE FROM clientes WHERE id=?", (cliente_id,))
            conn.commit()
            conn.close()
            self.cargar_clientes()
            messagebox.showinfo("Borrar", "Cliente borrado")

    def limpiar(self):
        self.nombre.set("")
        self.cedula.set("")
        messagebox.showinfo("Limpiar", "Campos limpiados")

    def generar_documento(self):
        doc = Document()
        doc.add_heading('Reclamación directa')
        doc.add_paragraph(f"Yo, {self.nombre.get()}, identificado con cédula {self.cedula.get()},")
        doc.add_paragraph("respetuosamente me permito presentar la siguiente reclamación en ejercicio de mi derecho al retracto, sobre el vuelo adquirido mediante tiquete adjunto. Lo anterior por cuanto no deseo hacer uso del servicio contratado y me encuentro dentro de la oportunidad para hacerlo. ")
        doc.add_paragraph("Argumentos jurídicos: Artículo 23 de la Constitucional, ley 1755 de 2015, artículo 47 del Estatuto del Consumidor y los RAC 3." )
        doc.add_paragraph("Pruebas: 1.	Documento identificación 2.	Tiquete")
        
        # Guardar el documento en el escritorio del usuario
        desktop_path = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
        doc_path = os.path.join(desktop_path, 'derecho_peticion.docx')
        doc.save(doc_path)
        messagebox.showinfo("Documento", f"Documento generado exitosamente en: {doc_path}")

if __name__ == "__main__":
    root = tk.Tk()
    app = CRUDCliente(root)
    root.mainloop()
